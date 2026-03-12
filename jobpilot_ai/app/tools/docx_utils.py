"""DOCX file helpers using python-docx.

Reusable across jobpilot_ai for reading, writing, and extracting text
from Word documents (.docx).
"""

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType


def read_docx(path: str | Path) -> DocumentType:
    """Load a .docx file and return a python-docx Document.

    Args:
        path: File path to the .docx file.

    Returns:
        A python-docx Document instance for further manipulation.

    Raises:
        FileNotFoundError: If the file does not exist.
        docx.opc.exceptions.PackageNotFoundError: If the file is not a valid .docx.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")
    return Document(path)


def write_docx(document: DocumentType, path: str | Path) -> Path:
    """Save a python-docx Document to a .docx file.

    Args:
        document: A python-docx Document instance (e.g. from read_docx).
        path: Output file path.

    Returns:
        The resolved output path (Path).

    Raises:
        OSError: If the file cannot be written (e.g. permission, disk).
    """
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return path.resolve()


def extract_text_from_docx(path: str | Path) -> str:
    """Extract plain text from a .docx file (paragraphs and tables).

    Args:
        path: File path to the .docx file.

    Returns:
        Concatenated text from all paragraphs and table cells, separated
        by newlines. No formatting is preserved.
    """
    doc = read_docx(path)
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n".join(parts)

